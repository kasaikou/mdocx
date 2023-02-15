import abc
import cairosvg
import docx
import docx.document
import docx.enum.text
import docx.shared
import docx.text.paragraph
import docx.text.run
import PIL.Image
import io
from typing import Optional

# interfaces


class MdocxText(metaclass=abc.ABCMeta):
    @abc.abstractmethod
    def push_to_paragraph(
        self,
        target: docx.text.paragraph.Paragraph,
    ) -> docx.text.paragraph.Run:
        raise NotImplementedError()


class MdocxParagraphContent(metaclass=abc.ABCMeta):
    @abc.abstractmethod
    def push_to_document(
        self,
        target: docx.document.Document,
    ) -> None:
        raise NotImplementedError


# util functions


def render_svg(svg: bytes, dpi: int = 220) -> bytes:
    return cairosvg.svg2png(bytestring=svg, dpi=dpi)


class MarkdownStatus:
    def __init__(self, line: int):
        self.line = line

    def dict(self) -> dict[str, str]:
        return {"line", str(self.line)}

    @staticmethod
    def from_dict(dict: dict, dictName: str):
        line = get_key("line", dict, dictName)
        return MarkdownStatus(line)


def get_key(
    key: str,
    dict: dict,
    dictName: Optional[str] = None,
    status: Optional[MarkdownStatus] = None,
):
    if key not in dict:
        dictName = dictName if dictName else "dict"
        status = f" ({status.dict()})" if status else ""
        raise RuntimeError(f"'{key}' not in {dictName}{status}")
    return dict[key]


# implements


class MdocxTextOptions:
    def __init__(self, bold: bool, italic: bool) -> None:
        self.bold = bold
        self.italic = italic


class MdocxReference:
    def __init__(
        self,
        key: str,
        status: MarkdownStatus,
        displayName: str,
    ):
        self.key = key
        self.status = status
        self.displayName = displayName

    @staticmethod
    def from_dict(dict: dict[str], dictName: str):
        status = get_key("status", dict, dictName)
        status = MarkdownStatus.from_dict(status)
        key = get_key("key", dict, dictName, status)
        displayName = get_key("displayName", dict, dictName, status)

        return MdocxReference(key, displayName, status)


ReferenceDict = dict[str, MdocxReference]


class NormalText(MdocxText):
    def __init__(self, expression: str, options: MdocxTextOptions):
        self.expression = expression
        self.options = options

    def push_to_paragraph(
        self, target: docx.text.paragraph.Paragraph
    ) -> docx.text.run.Run:
        run = target.add_run(text=self.expression)
        run.bold = self.options.bold
        run.italic = self.options.italic
        return run

    @staticmethod
    def from_dict(dict: dict[str], dictName: str, status: MarkdownStatus):
        expression = get_key("expression", dict, dictName, status)
        return NormalText(expression=expression)


class SvgText(MdocxText):
    def __init__(self, svg: str) -> None:
        self.png = render_svg(bytes(svg))

    def push_to_paragraph(
        self, target: docx.text.paragraph.Paragraph
    ) -> docx.text.run.Run:
        run = target.add_run()
        run.add_picture(io.BytesIO(self.png))
        return run

    @staticmethod
    def from_dict(dict: dict[str], dictName: str, status: MarkdownStatus):
        svg = get_key("svg", dict, dictName, status)
        return SvgText(svg)


class ReferenceText(MdocxText):
    def __init__(self, referenceKey: str):
        self.referenceKey = referenceKey

    def push_to_paragraph(
        self, target: docx.text.paragraph.Paragraph
    ) -> docx.text.run.Run:
        run = target.add_run(text=self.referenceKey)
        run.font.superscript = True
        return run

    @staticmethod
    def from_dict(
        dict: dict[str],
        dictName: str,
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ):
        key = get_key("key", dict, dictName, status)
        if key not in referenceDict:
            raise RuntimeError(f"unknown reference key '{key}' ({status.dict()})")

        return ReferenceText(referenceDict[key].displayName)


def push_mdocx_texts(
    texts: list[dict[str]],
    status: MarkdownStatus,
    referenceDict: ReferenceDict,
    target: docx.text.paragraph.Paragraph,
):
    for text in texts:
        textType = get_key("type", text, "text")

        match textType:
            case "normal":
                push = NormalText.from_dict(text, "normal text", status)

            case "svg":
                push = SvgText.from_dict(text, "svg text", status)

            case "reference":
                push = ReferenceText.from_dict(
                    text, "reference text", status, referenceDict
                )

            case _:
                raise RuntimeError(f"unknown text type '{textType}' ({status.dict()})")

        push.push_to_paragraph(target=target)


class NormalParagraphContent(MdocxParagraphContent):
    def __init__(
        self,
        texts: list[dict[str]],
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ) -> None:
        self.texts = texts
        self.status = status
        self.referenceDict = referenceDict

    def push_to_document(
        self,
        target: docx.document.Document,
    ) -> Optional[docx.text.paragraph.Paragraph]:
        paragraph = target.add_paragraph()
        push_mdocx_texts(
            texts=self.texts,
            status=self.status,
            referenceDict=self.referenceDict,
            target=paragraph,
        )

    @staticmethod
    def from_dict(
        dict: dict[str],
        dictName: str,
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ):
        texts = get_key("text", dict, dictName, status)
        return NormalParagraphContent(texts, status, referenceDict)


class ImageParagraphContent(MdocxParagraphContent):
    def __init__(
        self,
        filenames: list[str],
        description: list[dict[str]],
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ):
        self.filenames = filenames
        self.description = description
        self.status = status
        self.referenceDict = referenceDict

    def push_to_document(
        self,
        target: docx.document.Document,
    ) -> Optional[docx.text.paragraph.Paragraph]:
        # image size
        images = [PIL.Image.open(self.filenames[0])]
        imageBaseWidth, imageBaseHeight = images[0].size
        imageRate: list[float] = [float(imageBaseWidth)]
        imageTotalWidth = imageBaseWidth

        for filename in self.filenames[1:]:
            images.append(PIL.Image.open(filename))
            w, h = images[-1].size
            scaledWidth = w * (imageBaseHeight / h)
            imageRate.append(float(scaledWidth))
            imageTotalWidth += scaledWidth

        for i, _ in enumerate(imageRate):
            imageRate[i] /= imageTotalWidth

        # image
        paragraph = target.add_paragraph()
        paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        for i, filename in enumerate(self.filenames):
            paragraph.add_run().add_picture(
                filename, width=docx.shared.Inches(5 * imageRate[i])
            )

        # description
        paragraph = target.add_paragraph()
        paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        push_mdocx_texts(self.description, self.status, self.referenceDict, paragraph)

    @staticmethod
    def from_dict(
        dict: dict[str],
        dictName: str,
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ):
        filenames = get_key("filenames", dict, dictName, status)
        description = get_key("description", dict, dictName, status)
        return ImageParagraphContent(filenames, description, status, referenceDict)


class MathjaxSvgParagraphContent(MdocxParagraphContent):
    def __init__(self, svg: str):
        self.svg = svg

    def push_to_document(self, target: docx.document.Document) -> None:
        paragraph = target.add_paragraph()
        paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(io.BytesIO(self))

    @staticmethod
    def from_dict(
        dict: dict[str],
        dictName: str,
        status: MarkdownStatus,
    ):
        svg = get_key("svg", dict, dictName, status)
        return MathjaxSvgParagraphContent(svg)


class NewpageParagraphContent(MdocxParagraphContent):
    def __init__(self):
        pass

    def push_to_document(self, target: docx.document.Document) -> None:
        target.add_page_break()


class HeadingParagraphContent(MdocxParagraphContent):
    def __init__(
        self,
        level: int,
        text: list[dict[str]],
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ):
        self.level = level
        self.text = text
        self.status = status
        self.referenceDict = referenceDict

    def push_to_document(self, target: docx.document.Document) -> None:
        paragraph = target.add_paragraph("", f"Heading {self.level}")
        push_mdocx_texts(self.text, self.status, self.referenceDict, paragraph)

    @staticmethod
    def from_dict(
        dict: dict[str],
        dictName: str,
        status: MarkdownStatus,
        referenceDict: ReferenceDict,
    ):
        level = get_key("level", dict, dictName, status)
        texts = get_key("text", dict, dictName, status)
        return HeadingParagraphContent(level, texts, status, referenceDict)


def push_paragraphs(
    paragraphs: list[dict[str]],
    target: docx.document.Document,
    referenceDict: ReferenceDict,
):
    for paragraph in paragraphs:
        status = get_key("status", paragraph, "paragraph")
        status = MarkdownStatus.from_dict(status, "paragraph status")

        content = get_key("content", paragraph, "paragraph")
        contentType = get_key("type", content, "paragraph content")

        match contentType:
            case "normal":
                push = NormalParagraphContent.from_dict(
                    content, "normal paragraph content", status, referenceDict
                )

            case "image":
                push = ImageParagraphContent.from_dict(
                    content, "image paragraph content", status, referenceDict
                )

            case "mathjax":
                push = MathjaxSvgParagraphContent.from_dict(
                    content, "mathjax paragraph content", status
                )

            case "newpage":
                push = NewpageParagraphContent()

            case "heading":
                push = HeadingParagraphContent.from_dict(
                    content, "heading paragraph content", status, referenceDict
                )

            case _:
                raise RuntimeError(
                    f"unknown paragraph content type '{contentType}' ({status.dict()})"
                )

        push.push_to_document(target)


class Config:
    def __init__(
        self,
        title: str,
        description: Optional[str],
        author: Optional[str],
        styleTemplateFilename: Optional[str],
        destinationFilename: str,
    ):
        self.title = title
        self.description = description
        self.author = author
        self.styleTemplateFilename = styleTemplateFilename
        self.destinationFilename = destinationFilename

    def push_to_document(self, target: docx.document.Document):
        target.core_properties.title = self.title
        if self.author:
            target.core_properties.author = self.author
        if self.description:
            target.core_properties.subject = self.description
        if self.styleTemplateFilename:
            target.styles = docx.Document(self.styleTemplateFilename).styles

    def save(self, document: docx.document.Document):
        document.save(self.destinationFilename)

    @staticmethod
    def from_dict(self, dict: dict[str], dictName: str):
        title = get_key("title", dict, dictName)
        description = dict["description"] if "description" in dict else None
        author = dict["authorName"] if "authorName" in dict else None
        styleTemplateFilename = (
            dict["styleTemplateFilename"] if "styleTmplateFilename" in dict else None
        )
        destinationFilename = get_key("destination", dict, dictName)

        return Config(
            title=title,
            description=description,
            author=author,
            styleTemplateFileName=styleTemplateFilename,
            destinationFilename=destinationFilename,
        )


if __name__ == "__main__":
    import json, sys

    data = json.loads(sys.stdin)
    document = docx.Document()
    config = get_key("config", dict, "body")
    config = Config.from_dict(config)
    config.push_to_document(document)

    references = get_key("references", dict, "body")
    references: ReferenceDict = {
        key: MdocxReference.from_dict(reference, "reference")
        for key, reference in references.items()
    }
    paragraphs = get_key("paragraphs", dict, "body")
    push_paragraphs(paragraphs, document, references)

    config.save(document)
